import asyncio
import os
import re
import subprocess
import time
from urllib.parse import parse_qs, urlparse

import docx
import docx.opc.constants
import docx.oxml
import dotenv
import openai
import pandas as pd
import tiktoken
import yt_dlp as youtube_dl
from docx import Document
from docx.shared import Inches
from langdetect import detect
from PIL import Image
from pytube import YouTube

from .yt2t import YT2T

dotenv.load_dotenv(".env")
openai.api_key = os.environ.get("API_KEY")
encoding = tiktoken.get_encoding("cl100k_base")


# a.ru - автоматически сгенерированные английские
# ru - русский
# en - английский
def get_lang_clean_name(lang_name: str) -> str:
    """
    Получение чистого названия для языка
    """
    if "." in lang_name:
        return lang_name.split(".", 1)[1]
    else:
        return lang_name


def detect_language(text):
    lang = detect(text)
    if lang == "ru":
        return "ru-RU"
    elif lang == "eng":
        return "en-US"
    else:
        return None


def generate_subtitles(lang, yt=None, url=None):
    converter = YT2T()
    df = converter.url2text(urlpath=url, audioformat="flac", yt=yt, lang=lang)
    return df


def detect_lang_for_vid(dict_of_lang_subtitles, title):
    automatic_langs = {"a.ru": "ru-RU", "a.en": "en-US"}
    for lang in automatic_langs.keys():
        if lang in dict_of_lang_subtitles:
            return automatic_langs[lang]
    return detect_language(title)


def set_capital_and_remove_punctuation_marks(df: pd.DataFrame):
    prev_text = "##"
    for index, row in df.iterrows():
        text: str = row["text"]
        if len(prev_text) == 1:
            prev_text = text
            prev_text = text
            continue
        text = text.strip()
        if (
            prev_text == "##"
            or prev_text[-1] in [".", "?", "!"]
            and prev_text[-2] not in [".", "?", "!", ","]
        ):
            text = text.capitalize()

        prev_word = "##"
        df.at[index, "text"] = ""
        for word in text.split(" "):
            word = word.strip()
            word = word.strip()
            if len(prev_word) == 1:
                if prev_word != "##":
                    df.at[index, "text"] += " "
                if word[-1] == "." and word[-2] in ["?", "!"]:
                    word = word[:-1]
                df.at[index, "text"] += word
                prev_word = word
                continue
            if prev_word[-1] in [".", "?", "!"] and prev_word[-2] not in [
                ".",
                "?",
                "!",
                ",",
            ]:
                word = word.capitalize()

            if prev_word != "##":
                df.at[index, "text"] += " "

            if word[-1] == "." and word[-2] in ["?", "!"]:
                word = word[:-1]
            df.at[index, "text"] += word

            prev_word = word
        prev_text = df.at[index, "text"]
    return df


def concatenate_text(df):
    concatenated_text = ""

    # Проходимся по каждой строке датафрейма
    for index, row in df.iterrows():
        text = row["text"]

        # Удаляем все вхождения ".." из текста
        cleaned_text = text.replace("..", "")

        # Объединяем очищенный текст с предыдущими объединенными строками
        concatenated_text += cleaned_text

    # Возвращаем объединенный текст в виде строки
    return concatenated_text


def remove_rows_without_letters_and_numbers(df):
    # Создаем пустой список для хранения индексов строк, которые нужно удалить
    rows_to_remove = []

    # Проходимся по каждой строке датафрейма
    for index, row in df.iterrows():
        text = row["text"]

        # Проверяем, содержит ли поле "text" буквы, цифры или символы кириллицы
        if not re.search("[a-zA-Zа-яА-Я0-9]", text):
            rows_to_remove.append(index)

    # Удаляем строки из датафрейма по полученным индексам
    df = df.drop(rows_to_remove)

    return df


def create_annotation(str, limit_word):

    # 2.8 - среднее увеличение количества токенов по сравнение с количеством слов

    # Rate limit reached - error limit num of message
    # This model's maximum context length is 4097 tokens. - error extend len of message
    limit_tokens = round(limit_word * 2.8)
    if limit_tokens > 2600:
        limit_tokens = 2600

    encoding = tiktoken.encoding_for_model("gpt-3.5-turbo")
    tokens = encoding.encode(str)
    comp_str = encoding.decode(tokens)[:min(len(tokens), 4096 - limit_tokens)]

    message = f"Напиши аннотацию по данному тексту: {comp_str}. В ответе верни только аннотацию."
    response = None
    len_ext_error = "This model's maximum context length"
    lim_num_mes_error = "Rate limit reached"
    done = False
    while not done and limit_tokens > 0:
        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                temperature=0,
                max_tokens=limit_tokens,
                messages=[{"role": "user", "content": f"{message}"}],
            )
            done = True
        except openai.InvalidRequestError as e:
            if len_ext_error in e._message:
                limit_tokens -= 150
            if lim_num_mes_error in e._message:
                time.sleep(21)
            done = False
    # print response
    content_value = response["choices"][0]["message"]["content"]
    return content_value


def get_title(url):
    return YouTube(url).title


def get_subtitles_for_yt(link: str):
    """
    Получение субтитров для yt видео
    """

    try:
        yt = YouTube(link)
        dict_of_lang_subtitles = yt.captions
        title = yt.title
        lang_for_vid = detect_lang_for_vid(dict_of_lang_subtitles, title)

        # langs = ["ru", "en"]
        # for lang in langs:
        #     if lang in dict_of_lang_subtitles:
        #         subtitles = yt.captions[lang].generate_srt_captions()
        #         df = parse_subtitles(subtitles)
        #         break
        #
        # else:
        #     if lang_for_vid is not None:
        #         df = generate_subtitles(lang=lang_for_vid, yt=yt)

        df = generate_subtitles(lang=lang_for_vid, yt=yt)

        df = remove_rows_without_letters_and_numbers(df)

        df = set_capital_and_remove_punctuation_marks(df)

        return df

    except Exception as e:
        print("Произошла ошибка:", e)


def parse_subtitles(subtitles_string):
    lines = subtitles_string.strip().split("\n\n")
    data = []

    for line in lines:
        parts = line.split("\n")
        index = parts[0]

        time_regex = re.compile(
            r"(\d{2}:\d{2}:\d{2},\d{3})\s*-->\s*(\d{2}:\d{2}:\d{2},\d{3})"
        )
        time_match = time_regex.search(parts[1])
        start_time = time_match.group(1)
        end_time = time_match.group(2)

        text = " ".join(parts[2:])

        data.append((text, start_time, end_time))

    df = pd.DataFrame(data, columns=["text", "start_time", "end_time"])
    return df


def extract_picture_from_yt_video(
    url: str, start_time: str = "00:00:00.000", nm_pct_with_ext: str = "output.jpg"
):
    ydl_opts = {
        "quiet": True,  # Отключение вывода информации от youtube_dl
    }
    with youtube_dl.YoutubeDL(ydl_opts) as ydl:
        info_dict = ydl.extract_info(url, download=False)
        formats = info_dict.get("formats", [])
        mp4_formats = [format for format in formats if format.get("ext") == "mp4"]
        max_quality = max(mp4_formats, key=lambda x: x["quality"])
        video_url = max_quality["url"]

    if video_url:
        ffmpeg_command = f'ffmpeg -ss {start_time} -i "{video_url}" -frames:v 1 -update 1 -y {nm_pct_with_ext}'
        subprocess.run(ffmpeg_command, shell=True)


def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(
        docx.oxml.shared.qn("r:id"),
        r_id,
    )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement("w:r")

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement("w:rPr")

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement("w:color")
        c.set(docx.oxml.shared.qn("w:val"), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement("w:u")
        u.set(docx.oxml.shared.qn("w:val"), "none")
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def get_seconds(time_str) -> int:
    # Разделение строки на составляющие
    hours, minutes, seconds = time_str.split(":")
    # Извлечение значения секунд и преобразование в целое число
    seconds = int(seconds.split(".")[0])

    return seconds


def delete_file(path: str) -> bool:
    if os.path.exists(path):
        os.remove(path)
        return True
    else:
        return False


def get_yt_vid_id(url: str) -> str:
    url_data = urlparse(url)
    query = parse_qs(url_data.query)
    if "v" in query.keys():
        video_id = query["v"][0]
    else:
        video_id = url_data.geturl().split("/")[-1]
    return video_id


def create_doc(
    df: pd.DataFrame,
    url: str,
    word_limit_annotation: int = 1000,
    add_annonation: bool = True,
    add_name: str = "",
):
    # Создание нового документа
    video_id = get_yt_vid_id(url)
    name_of_doc_file = "data/docx_file/" + video_id + add_name + ".docx"
    image_path = "data/images/" + video_id + add_name + "_image.png"
    temp_image_path = "data/images/" + video_id + add_name + "_temp_image.png"

    title = get_title(url)

    doc = Document()
    # Добавление текстового содержимого из датафрейма в документ
    doc.add_heading(f"{title}", level=1)
    if add_annonation is True:
        annonation = create_annotation(concatenate_text(df), word_limit_annotation)
        doc.add_paragraph(annonation)
        doc.add_page_break()
    num_of_paragraph = 0
    for index, row in df.iterrows():
        if row["text"].strip()[0].isupper():
            num_of_paragraph += 1
            doc.add_heading(f"Параграф {num_of_paragraph}", level=2)

        p = doc.add_paragraph("")

        time_code = get_seconds(row["start_time"])
        link = url + f"&t={time_code}"

        add_hyperlink(p, link, row["start_time"], "FF8822", True)

        doc.add_paragraph(row["text"])
        # Добавление изображений в документ

        extract_picture_from_yt_video(
            url, start_time=row["start_time"], nm_pct_with_ext=image_path
        )

        img = Image.open(image_path)

        # Определение размеров изображения в дюймах (пропорционально)
        img_width, img_height = img.size
        aspect_ratio = img_width / img_height
        desired_width = Inches(6)
        desired_height = desired_width / aspect_ratio

        # Масштабирование изображения с сохранением пропорций
        img.thumbnail((desired_width, desired_height))

        # Сохранение временной копии масштабированного изображения в формате JPEG

        img.save(temp_image_path, "JPEG")
        # Добавление изображения в документ
        doc.add_picture(temp_image_path, width=desired_width, height=desired_height)

        # Разделитель между разделами документа
        # doc.add_page_break()

    delete_file(image_path)
    delete_file(temp_image_path)
    # Сохранение документа
    doc.save(name_of_doc_file)
    if add_annonation is True:
        return name_of_doc_file, annonation
    else:
        return name_of_doc_file


def get_doc_from_url(url: str, word_limit_annotation: int = 1000):
    try:
        df = get_subtitles_for_yt(url)
        # video_id = get_yt_vid_id(url)
        # path = "data/subtitle/" + video_id + ".csv"
        # df.to_csv(path)
        name_of_doc_file, annonation = create_doc(df, url, word_limit_annotation)
        return name_of_doc_file, annonation, df
    except Exception as e:
        print("Произошла ошибка:", e)
        return None, None, None


def form_paragraph_for_gen(df: pd.DataFrame):
    df["text"].iloc[-1] = df["text"].iloc[-1].replace("..", "").replace(",.", "")
    paragraph_df = pd.DataFrame({"text": [], "start_time": [], "end_time": []})
    limit_symbols = 150
    union_row = ""
    union_start_time = None
    union_by_len_row = ""

    for index, row in df.iterrows():
        if ".." in row["text"] or ",." in row["text"]:
            if union_row == "":
                union_row = row["text"]
                union_start_time = row["start_time"]
            else:
                union_row += " " + row["text"]
            union_end_time = row["end_time"]
        else:
            if union_row != "":
                union_row += " " + row["text"]
                union_end_time = row["end_time"]
                if union_by_len_row != "":
                    union_by_len_row += " "
                union_by_len_row += union_row
                union_by_len_start_time = union_start_time
                union_by_len_end_time = union_end_time
                union_row = ""
                union_start_time = None
                union_end_time = None

            if union_by_len_row == "":
                union_by_len_row = row["text"]
                union_by_len_start_time = row["start_time"]
                union_by_len_end_time = row["end_time"]
            if len(union_by_len_row) >= limit_symbols:
                new_row = pd.DataFrame(
                    {
                        "text": [union_by_len_row],
                        "start_time": [union_by_len_start_time],
                        "end_time": [union_by_len_end_time],
                    }
                )
                paragraph_df = pd.concat([paragraph_df, new_row], ignore_index=True)
                union_by_len_row = ""
                union_by_len_start_time = None
                union_by_len_end_time = None
    if union_by_len_row != "":
        new_row = pd.DataFrame(
            {
                "text": [union_by_len_row],
                "start_time": [union_by_len_start_time],
                "end_time": [union_by_len_end_time],
            }
        )
        paragraph_df = pd.concat([paragraph_df, new_row], ignore_index=True)

    return paragraph_df


def gen_text_based_on_paragraph(
    df_subtitle: pd.DataFrame, limit_article_length: int, url: str
):
    df_form_paragraph = form_paragraph_for_gen(df_subtitle)
    num_of_paragraph = df_form_paragraph.shape[0]
    len_of_one_paragraph = round(limit_article_length / num_of_paragraph)
    # 2.8 - среднее увеличение количества токенов по сравнение с количеством слов

    # Rate limit reached - error limit num of message
    # This model's maximum context length is 4097 tokens. - error extend len of message
    limit_tokens = round(len_of_one_paragraph * 2.8)
    if limit_tokens > 2500:
        limit_tokens = 2500
    response = None
    len_ext_error = "This model's maximum context length"
    lim_num_mes_error = "Rate limit reached"

    for index, row in df_form_paragraph.iterrows():
        response = None
        limit_tokens_row = limit_tokens

        encoding = tiktoken.encoding_for_model("gpt-3.5-turbo")
        tokens = encoding.encode(row['text'])
        comp_str = encoding.decode(tokens)[:min(len(tokens), 4096 - limit_tokens_row)]

        message = f"Cформируй связный красивый текст из данного текста: {comp_str}. В ответе верни только сам текст."
        done = False
        while not done and limit_tokens_row > 0:
            time.sleep(2)
            try:
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    temperature=0,
                    max_tokens=limit_tokens_row,
                    messages=[{"role": "user", "content": f"{message}"}],
                )
                done = True
            except openai.InvalidRequestError as e:
                if len_ext_error in e._message:
                    limit_tokens_row -= 150
                if lim_num_mes_error in e._message:
                    time.sleep(21)
                done = False
            except Exception as e:
                done = False
                time.sleep(21)
        # print response
        content_value = response["choices"][0]["message"]["content"]
        row["text"] = content_value

    name_of_doc_file = create_doc(
        df_form_paragraph, url, 0, False, add_name="_gen_vers_"
    )
    return name_of_doc_file


def get_all_articles(url, word_limit_annotation=1000, limit_article_length=100000):
    name_of_doc_file, annonation, df_subtitle = get_doc_from_url(
        url, word_limit_annotation=word_limit_annotation
    )
    name_of_doc_gen_file = gen_text_based_on_paragraph(
        df_subtitle, limit_article_length, url
    )
    return name_of_doc_file, name_of_doc_gen_file, annonation


# url = "https://www.youtube.com/watch?v=V6G3sPbgubY"
# print(get_all_articles(url))
